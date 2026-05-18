import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def add_blockquote(doc, text, bold=False):
    """添加引用块样式段落"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1.5)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.italic = True
    if bold:
        run.bold = True
    return para


def add_bold_text(doc, text, font_size=10.5, color=None, alignment=None):
    """添加加粗段落"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    if alignment:
        para.alignment = alignment
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = True
    if color:
        run.font.color.rgb = color
    return para


def add_normal_text(doc, text, font_size=10.5, color=None, spacing_after=4):
    """添加普通段落"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(spacing_after)
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    return para


def add_horizontal_rule(doc):
    """添加分隔线"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)
    # 用三个星号作为分隔
    run = para.add_run("✦ ✦ ✦")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(180, 180, 180)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para


def parse_inline(text):
    """解析行内格式：粗体、代码块、斜体"""
    parts = []
    i = 0
    while i < len(text):
        # 粗体 **...**
        if text[i:i+2] == '**' and i+2 < len(text):
            end = text.find('**', i+2)
            if end != -1:
                parts.append(('bold', text[i+2:end]))
                i = end + 2
                continue
        # 代码 `...`
        if text[i] == '`' and i+1 < len(text):
            end = text.find('`', i+1)
            if end != -1:
                parts.append(('code', text[i+1:end]))
                i = end + 1
                continue
        # 加粗单星号 *...*
        if text[i] == '*' and i+1 < len(text) and text[i:i+2] != '**':
            end = text.find('*', i+1)
            if end != -1:
                parts.append(('bold', text[i+1:end]))
                i = end + 1
                continue
        # 🔹 表情符号
        if text[i:i+2] in ('🔹', '🎤', '📝', '🎬', '⚠️', '📦', '📚'):
            parts.append(('text', text[i:i+2]))
            i += 2
            continue
        parts.append(('text', text[i]))
        i += 1
    return parts


def add_rich_paragraph(doc, text, font_size=10.5, indent_left=None, is_quote=False, spacing_after=4):
    """添加带格式的段落（支持粗体、代码等）"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(spacing_after)
    if indent_left:
        para.paragraph_format.left_indent = Cm(indent_left)

    if is_quote:
        # 引用风格：灰色斜体
        run = para.add_run(text)
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.italic = True
        return para

    parts = parse_inline(text)
    for part_type, part_text in parts:
        run = para.add_run(part_text)
        run.font.size = Pt(font_size)
        if part_type == 'bold':
            run.bold = True
        elif part_type == 'code':
            run.font.name = 'Consolas'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            run.font.color.rgb = RGBColor(200, 60, 60)
            run.font.size = Pt(9.5)
    return para


def add_list_item(doc, text, font_size=10.5, bullet="-"):
    """添加列表项"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)

    parts = parse_inline(text)
    run = para.add_run(f"{bullet} ")
    run.font.size = Pt(font_size)

    for part_type, part_text in parts:
        run = para.add_run(part_text)
        run.font.size = Pt(font_size)
        if part_type == 'bold':
            run.bold = True
        elif part_type == 'code':
            run.font.name = 'Consolas'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            run.font.color.rgb = RGBColor(200, 60, 60)
            run.font.size = Pt(9.5)
    return para


def convert_md_to_docx(md_path, docx_path):
    """将 Markdown 文件转换为 Word 文档"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # 跳过空行
        if not line.strip():
            i += 1
            continue

        # 水平分割线 ***
        if line.strip() == '***':
            add_horizontal_rule(doc)
            i += 1
            continue

        # 标题 # xxx
        if line.startswith('# ') and not line.startswith('## '):
            title_text = line[2:].strip()
            # 移除可能的 markdown 粗体标记
            title_text = title_text.replace('**', '')
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(16)
            run = para.add_run(title_text)
            run.font.size = Pt(18)
            run.bold = True
            run.font.color.rgb = RGBColor(30, 30, 30)
            i += 1
            continue

        # 引用 > xxx
        if line.startswith('> '):
            quote_text = line[2:].strip()
            # 移除 markdown 格式
            quote_text = quote_text.replace('**', '')
            add_rich_paragraph(doc, quote_text, font_size=10, is_quote=True)
            i += 1
            continue

        # 列表项 - xxx
        if line.strip().startswith('- '):
            item_text = line.strip()[2:]
            add_list_item(doc, item_text)
            i += 1
            continue

        # 带有 ** 开头的粗体标题行（如 **🎤 语音生成**）
        if line.strip().startswith('**') and not line.strip().startswith('**🎤') and not line.strip().startswith('**📝') and not line.strip().startswith('**🎬') and not line.strip().startswith('**⚠️') and not line.strip().startswith('**📦') and not line.strip().startswith('**📚'):
            text = line.strip()
            # 粗体段落
            text_clean = text.replace('**', '')
            add_bold_text(doc, text_clean, font_size=11, color=RGBColor(30, 30, 30))
            i += 1
            continue

        # 带有 emoji 的粗体小标题
        if line.strip().startswith('**🎤') or line.strip().startswith('**📝') or line.strip().startswith('**🎬') or line.strip().startswith('**⚠️') or line.strip().startswith('**📦') or line.strip().startswith('**📚') or line.strip().startswith('**🔹'):
            text = line.strip().replace('**', '')
            add_bold_text(doc, text, font_size=12, color=RGBColor(40, 40, 40))
            i += 1
            continue

        # 普通段落
        text = line.strip()
        if text:
            add_rich_paragraph(doc, text)

        i += 1

    doc.save(docx_path)
    print(f"已导出到: {docx_path}")


if __name__ == '__main__':
    import sys

    md_file = sys.argv[1] if len(sys.argv) > 1 else r"f:\1_Project\Gitter\docs\宣传文案\202605171126-2.md"
    docx_file = sys.argv[2] if len(sys.argv) > 2 else r"f:\1_Project\Gitter\docs\宣传文案\Gitter_宣传文案.docx"

    convert_md_to_docx(md_file, docx_file)
